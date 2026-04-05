"""Извлечение текста из учебных файлов (DOCX и др.)."""
import io
import logging

logger = logging.getLogger("exam_system.document_extract")


def text_from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = " ".join((c.text or "").strip() for c in row.cells if (c.text or "").strip())
            if cells:
                parts.append(cells)
    return "\n".join(parts)


def text_from_docx_file(path: str) -> str:
    with open(path, "rb") as f:
        return text_from_docx(f.read())
